from fastapi import FastAPI, Request, Form, UploadFile, File
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel
from typing import Dict, Any, List
from io import BytesIO
from docx import Document
from datetime import datetime
import json
import os
import tempfile
from docx2pdf import convert as docx2pdf_convert
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import xml.etree.ElementTree as ET
import urllib.request


class RequestData(BaseModel):
    # Kept for potential future validation; dynamic form parsing is used now
    applicant_name: str = ""
    applicant_address: str = ""
    parcel_numbers: str = ""
    precinct: str = ""
    municipality: str = ""
    investment_description: str = ""
    land_use: str = ""
    building_height: str = ""
    building_area: str = ""
    parking: str = ""


def load_municipality_templates() -> Dict[str, Dict[str, str]]:
    config_path = os.path.join(os.path.dirname(__file__), "municipalities.json")
    if not os.path.exists(config_path):
        # Basic default if config missing
        return {
            "konopnica": {
                "name": "Konopnica",
                "templates": {
                    "analysis": "templates/konopnica_analysis.xml",
                    "decision": "templates/konopnica_decision.xml"
                },
                "header": "Analiza urbanistyczna - Gmina Konopnica",
                "intro": "Analiza funkcji oraz cech zabudowy i zagospodarowania terenu w zakresie warunków, o których mowa w art. 61 ust. 1–6 ustawy o planowaniu i zagospodarowaniu przestrzennym.",
                "footer": "Urząd Gminy Konopnica"
            }
        }
    with open(config_path, "r", encoding="utf-8") as f:
        return json.load(f)


def compare_values(left: Dict[str, str], right: Dict[str, str]) -> Dict[str, Dict[str, Any]]:
    fields = left.keys() | right.keys()
    result: Dict[str, Dict[str, Any]] = {}
    for field in fields:
        lval = (left.get(field) or "").strip()
        rval = (right.get(field) or "").strip()
        result[field] = {
            "left": lval,
            "right": rval,
            "match": (lval.lower() == rval.lower())
        }
    return result


def sanitize_case_number(case_number: str) -> str:
    if not case_number:
        return ""
    return case_number.replace('.', '_').strip()


def generate_docx(analysis: Dict[str, str], gmina: str, templates_cfg: Dict[str, Dict[str, str]]) -> bytes:
    doc = Document()
    tpl = templates_cfg.get(gmina) or list(templates_cfg.values())[0]

    doc.add_heading(tpl.get("header", "Analiza urbanistyczna"), level=1)
    doc.add_paragraph(tpl.get("intro", ""))
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Pole"
    hdr_cells[1].text = "Wartość analizy"

    human_labels = field_labels()
    for key, label in human_labels.items():
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = (analysis.get(key) or "").strip()

    doc.add_paragraph("")
    # Optional narrative sections
    wyniki = (analysis.get("wyniki_analizy") or "").strip()
    if wyniki:
        doc.add_heading("Wyniki analizy", level=2)
        doc.add_paragraph(wyniki)

    uzas = (analysis.get("uzasadnienie") or "").strip()
    if uzas:
        doc.add_heading("Uzasadnienie", level=2)
        doc.add_paragraph(uzas)

    podstawy = (analysis.get("podstawy_prawne") or "").strip()
    if podstawy:
        doc.add_heading("Podstawy prawne", level=2)
        doc.add_paragraph(podstawy)

    doc.add_paragraph(f"Gmina: {gmina}")
    doc.add_paragraph(tpl.get("footer", ""))
    doc.add_paragraph(datetime.now().strftime("Data generacji: %Y-%m-%d %H:%M"))

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


def _fill_placeholders(text: str, ctx: Dict[str, str]) -> str:
    out = text or ""
    for k, v in ctx.items():
        out = out.replace(f"{{{{{k}}}}}", str(v or ""))
    return out


def generate_docx_from_xml(analysis: Dict[str, str], wniosek: Dict[str, str], gmina: str, case_number: str) -> bytes:
    # Znajdź konfigurację gminy
    municipalities = load_municipality_templates()
    gmina_key = gmina.lower().replace("gmina ", "").replace(" ", "_")
    
    if gmina_key not in municipalities:
        gmina_key = "konopnica"  # fallback
    
    municipality_config = municipalities[gmina_key]
    xml_path = os.path.join(os.path.dirname(__file__), municipality_config["templates"]["analysis"])
    
    if not os.path.exists(xml_path):
        # Fallback to standard docx if xml missing
        return generate_docx(analysis, gmina, load_municipality_templates())

    # Merge analysis and wniosek data for context
    ctx = {}
    ctx.update(wniosek)  # Zawiera pola z prefixem wniosek_
    ctx.update(analysis)  # Zawiera pola bez prefixu
    ctx["gmina"] = f"Gmina {gmina}"
    ctx["case_number"] = case_number
    doc = Document()
    tree = ET.parse(xml_path)
    root = tree.getroot()

    header = root.find("Header")
    if header is not None:
        title = _fill_placeholders((header.findtext("Title") or ""), ctx)
        if title:
            doc.add_heading(title, level=1)
        subtitle = _fill_placeholders((header.findtext("Subtitle") or ""), ctx)
        if subtitle:
            p = doc.add_paragraph()
            run = p.add_run(subtitle)
            run.bold = True
        legal = _fill_placeholders((header.findtext("LegalBase") or ""), ctx)
        if legal:
            doc.add_paragraph(legal)

    case_el = root.find("Case")
    if case_el is not None:
        nr = _fill_placeholders((case_el.findtext("CaseNumber") or ""), ctx)
        if nr:
            para = doc.add_paragraph()
            run = para.add_run(f"Numer sprawy: {nr}")
            run.bold = True
        plots = _fill_placeholders((case_el.findtext("Plots") or ""), ctx)
        if plots:
            doc.add_paragraph(plots)

    for section in root.findall("Section"):
        title = section.get("title") or ""
        title = _fill_placeholders(title, ctx)
        if title:
            doc.add_heading(title, level=2)
        for child in section:
            if child.tag == "Paragraph":
                doc.add_paragraph(_fill_placeholders(child.text or "", ctx))
            elif child.tag == "Point":
                idx = child.get("index") or ""
                t = child.get("title") or ""
                para = doc.add_paragraph()
                run = para.add_run(f"{idx}. { _fill_placeholders(t, ctx) }")
                run.bold = True
                for sp in child:
                    if sp.tag == "Text":
                        doc.add_paragraph(_fill_placeholders(sp.text or "", ctx))
                    elif sp.tag == "Subpoint":
                        sidx = sp.get("index") or ""
                        stitle = _fill_placeholders(sp.get("title") or "", ctx)
                        pr = doc.add_paragraph()
                        r = pr.add_run(f"{sidx}) {stitle}")
                        r.bold = True
                        for sub in sp:
                            if sub.tag == "Text":
                                doc.add_paragraph(_fill_placeholders(sub.text or "", ctx))
                            elif sub.tag == "List":
                                for item in sub.findall("Item"):
                                    label = _fill_placeholders(item.get("label") or "", ctx)
                                    val = _fill_placeholders(item.text or "", ctx)
                                    doc.add_paragraph(f"- {label}: {val}")

    annex = root.find("Annex")
    if annex is not None:
        doc.add_heading(_fill_placeholders(annex.get("title") or "Załączniki", ctx), level=2)
        for ch in annex:
            if ch.tag == "Text":
                doc.add_paragraph(_fill_placeholders(ch.text or "", ctx))

    footer = root.find("Footer")
    if footer is not None:
        sign_hint = footer.findtext("SignHint") or ""
        doc.add_paragraph("")
        doc.add_paragraph(_fill_placeholders(sign_hint, ctx))

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


def generate_pdf_from_docx_bytes(docx_bytes: bytes) -> bytes:
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "doc.docx")
        pdf_path = os.path.join(tmpdir, "doc.pdf")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)
        # Uses MS Word on Windows; requires Word installed
        docx2pdf_convert(docx_path, pdf_path)
        with open(pdf_path, "rb") as f:
            return f.read()


def generate_pdf_basic(analysis: Dict[str, str], gmina: str, templates_cfg: Dict[str, Dict[str, str]]) -> bytes:
    bio = BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    # Ensure Polish diacritics: register DejaVu family (regular + bold) if available
    family = ensure_font_installed()
    font_regular_path = os.path.join(BASE_DIR, 'static', 'fonts', 'DejaVuSans.ttf')
    font_bold_path = os.path.join(BASE_DIR, 'static', 'fonts', 'DejaVuSans-Bold.ttf')
    base_font = 'Helvetica'
    if family and os.path.exists(font_regular_path):
        try:
            pdfmetrics.registerFont(TTFont('DejaVuSans', font_regular_path))
            if os.path.exists(font_bold_path):
                pdfmetrics.registerFont(TTFont('DejaVuSans-Bold', font_bold_path))
                pdfmetrics.registerFontFamily('DejaVu', normal='DejaVuSans', bold='DejaVuSans-Bold', italic='DejaVuSans', boldItalic='DejaVuSans-Bold')
            base_font = 'DejaVu'
        except Exception:
            base_font = 'Helvetica'
    # Apply explicit fonts
    if base_font == 'DejaVu':
        styles['Normal'].fontName = 'DejaVuSans'
        if 'Title' in styles.byName:
            styles['Title'].fontName = 'DejaVuSans-Bold'
        if 'Heading2' in styles.byName:
            styles['Heading2'].fontName = 'DejaVuSans-Bold'
    tpl = templates_cfg.get(gmina) or list(templates_cfg.values())[0]

    elements = []
    elements.append(Paragraph(tpl.get("header", "Analiza urbanistyczna"), styles['Title']))
    intro = tpl.get("intro", "")
    if intro:
        elements.append(Paragraph(intro, styles['Normal']))
        elements.append(Spacer(1, 12))

    labels = field_labels()
    data = [["Pole", "Wartość analizy"]]
    for key, label in labels.items():
        data.append([label, (analysis.get(key) or "").strip()])

    table = Table(data, colWidths=[200, 320])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('FONTNAME', (0,0), (-1,-1), base_font if base_font != 'DejaVu' else 'DejaVuSans'),
        ('FONTNAME', (0,0), (-1,0), 'DejaVuSans-Bold' if base_font == 'DejaVu' else 'Helvetica-Bold'),
    ]))
    elements.append(table)
    elements.append(Spacer(1, 12))
    # Narrative sections
    wyniki = (analysis.get("wyniki_analizy") or "").strip()
    if wyniki:
        elements.append(Paragraph("Wyniki analizy", styles['Heading2']))
        elements.append(Paragraph(wyniki.replace('\n', '<br/>'), styles['Normal']))
        elements.append(Spacer(1, 8))

    uzas = (analysis.get("uzasadnienie") or "").strip()
    if uzas:
        elements.append(Paragraph("Uzasadnienie", styles['Heading2']))
        elements.append(Paragraph(uzas.replace('\n', '<br/>'), styles['Normal']))
        elements.append(Spacer(1, 8))

    podstawy = (analysis.get("podstawy_prawne") or "").strip()
    if podstawy:
        elements.append(Paragraph("Podstawy prawne", styles['Heading2']))
        elements.append(Paragraph(podstawy.replace('\n', '<br/>'), styles['Normal']))
        elements.append(Spacer(1, 8))

    elements.append(Paragraph(f"Gmina: {gmina}", styles['Normal']))
    footer = tpl.get("footer", "")
    if footer:
        elements.append(Paragraph(footer, styles['Normal']))

    doc.build(elements)
    bio.seek(0)
    return bio.read()


def generate_pdf_from_xml(analysis: Dict[str, str], wniosek: Dict[str, str], gmina: str, case_number: str) -> bytes:
    # Znajdź konfigurację gminy
    municipalities = load_municipality_templates()
    gmina_key = gmina.lower().replace("gmina ", "").replace(" ", "_")
    
    if gmina_key not in municipalities:
        gmina_key = "konopnica"  # fallback
    
    municipality_config = municipalities[gmina_key]
    xml_path = os.path.join(os.path.dirname(__file__), municipality_config["templates"]["analysis"])
    
    if not os.path.exists(xml_path):
        return generate_pdf_basic(analysis, gmina, load_municipality_templates())

    # Merge analysis and wniosek data for context
    ctx = {}
    ctx.update(wniosek)  # Zawiera pola z prefixem wniosek_
    ctx.update(analysis)  # Zawiera pola bez prefixu
    ctx["gmina"] = f"Gmina {gmina}"
    ctx["case_number"] = case_number
    bio = BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    # Ensure Polish diacritics via DejaVu family
    family = ensure_font_installed()
    font_regular_path = os.path.join(BASE_DIR, 'static', 'fonts', 'DejaVuSans.ttf')
    font_bold_path = os.path.join(BASE_DIR, 'static', 'fonts', 'DejaVuSans-Bold.ttf')
    base_font = 'Helvetica'
    if family and os.path.exists(font_regular_path):
        try:
            pdfmetrics.registerFont(TTFont('DejaVuSans', font_regular_path))
            if os.path.exists(font_bold_path):
                pdfmetrics.registerFont(TTFont('DejaVuSans-Bold', font_bold_path))
                pdfmetrics.registerFontFamily('DejaVu', normal='DejaVuSans', bold='DejaVuSans-Bold', italic='DejaVuSans', boldItalic='DejaVuSans-Bold')
            base_font = 'DejaVu'
        except Exception:
            base_font = 'Helvetica'
    if base_font == 'DejaVu':
        styles['Normal'].fontName = 'DejaVuSans'
        if 'Title' in styles.byName:
            styles['Title'].fontName = 'DejaVuSans-Bold'
        if 'Heading2' in styles.byName:
            styles['Heading2'].fontName = 'DejaVuSans-Bold'
    elements = []

    tree = ET.parse(xml_path)
    root = tree.getroot()

    header = root.find("Header")
    if header is not None:
        title = _fill_placeholders((header.findtext("Title") or ""), ctx)
        if title:
            elements.append(Paragraph(title, styles['Title']))
        subtitle = _fill_placeholders((header.findtext("Subtitle") or ""), ctx)
        if subtitle:
            elements.append(Paragraph(f"<b>{subtitle}</b>", styles['Normal']))
        legal = _fill_placeholders((header.findtext("LegalBase") or ""), ctx)
        if legal:
            elements.append(Paragraph(legal, styles['Normal']))
        elements.append(Spacer(1, 8))

    case_el = root.find("Case")
    if case_el is not None:
        nr = _fill_placeholders((case_el.findtext("CaseNumber") or ""), ctx)
        if nr:
            elements.append(Paragraph(f"<b>Numer sprawy:</b> {nr}", styles['Normal']))
        plots = _fill_placeholders((case_el.findtext("Plots") or ""), ctx)
        if plots:
            elements.append(Paragraph(plots, styles['Normal']))
        elements.append(Spacer(1, 8))

    for section in root.findall("Section"):
        title = _fill_placeholders(section.get("title") or "", ctx)
        if title:
            elements.append(Paragraph(title, styles['Heading2']))
        for child in section:
            if child.tag == "Paragraph":
                elements.append(Paragraph(_fill_placeholders(child.text or "", ctx), styles['Normal']))
            elif child.tag == "Point":
                idx = child.get("index") or ""
                t = _fill_placeholders(child.get("title") or "", ctx)
                elements.append(Paragraph(f"<b>{idx}. {t}</b>", styles['Normal']))
                for sp in child:
                    if sp.tag == "Text":
                        elements.append(Paragraph(_fill_placeholders(sp.text or "", ctx), styles['Normal']))
                    elif sp.tag == "Subpoint":
                        sidx = sp.get("index") or ""
                        stitle = _fill_placeholders(sp.get("title") or "", ctx)
                        elements.append(Paragraph(f"<b>{sidx}) {stitle}</b>", styles['Normal']))
                        for sub in sp:
                            if sub.tag == "Text":
                                elements.append(Paragraph(_fill_placeholders(sub.text or "", ctx), styles['Normal']))
                            elif sub.tag == "List":
                                for item in sub.findall("Item"):
                                    label = _fill_placeholders(item.get("label") or "", ctx)
                                    val = _fill_placeholders(item.text or "", ctx)
                                    elements.append(Paragraph(f"- <b>{label}:</b> {val}", styles['Normal']))
        elements.append(Spacer(1, 6))

    annex = root.find("Annex")
    if annex is not None:
        elements.append(Paragraph(_fill_placeholders(annex.get("title") or "Załączniki", ctx), styles['Heading2']))
        for ch in annex:
            if ch.tag == "Text":
                elements.append(Paragraph(_fill_placeholders(ch.text or "", ctx), styles['Normal']))

    footer = root.find("Footer")
    if footer is not None:
        sign_hint = footer.findtext("SignHint") or ""
        if sign_hint:
            elements.append(Spacer(1, 12))
            elements.append(Paragraph(sign_hint, styles['Normal']))

    doc.build(elements)
    bio.seek(0)
    return bio.read()


def generate_decision_docx_from_xml(analysis: Dict[str, str], wniosek: Dict[str, str], gmina: str, case_number: str) -> bytes:
    """Generate DOCX decision document from XML template"""
    # Znajdź konfigurację gminy
    municipalities = load_municipality_templates()
    gmina_key = gmina.lower().replace("gmina ", "").replace(" ", "_")
    
    if gmina_key not in municipalities:
        gmina_key = "konopnica"  # fallback
    
    municipality_config = municipalities[gmina_key]
    xml_path = os.path.join(os.path.dirname(__file__), municipality_config["templates"]["decision"])
    
    if not os.path.exists(xml_path):
        raise FileNotFoundError(f"Decision XML template not found: {xml_path}")

    # Merge analysis and wniosek data for context
    # Używamy danych bezpośrednio - XML oczekuje pól z prefixem wniosek_
    ctx = {}
    ctx.update(wniosek)  # Zawiera pola z prefixem wniosek_
    ctx.update(analysis)  # Zawiera pola bez prefixu
    ctx["gmina"] = f"Gmina {gmina}"
    ctx["case_number"] = case_number
    # Fill missing date fields if not provided
    if "data" not in ctx or not ctx["data"]:
        ctx["data"] = datetime.now().strftime("%d.%m.%Y r.")
    if "data_wniosku" not in ctx or not ctx["data_wniosku"]:
        ctx["data_wniosku"] = datetime.now().strftime("%d.%m.%Y")
    if "data_uzupełnienia" not in ctx or not ctx["data_uzupełnienia"]:
        ctx["data_uzupełnienia"] = datetime.now().strftime("%d.%m.%Y")
    if "rodzaj_zabudowy" not in ctx or not ctx["rodzaj_zabudowy"]:
        ctx["rodzaj_zabudowy"] = "zabudowa mieszkaniowa jednorodzinna"

    doc = Document()
    tree = ET.parse(xml_path)
    root = tree.getroot()

    header = root.find("Header")
    if header is not None:
        ref_num = _fill_placeholders((header.findtext("ReferenceNumber") or ""), ctx)
        if ref_num:
            para = doc.add_paragraph()
            run = para.add_run(ref_num)
            run.bold = True
        place_date = _fill_placeholders((header.findtext("PlaceDate") or ""), ctx)
        if place_date:
            para = doc.add_paragraph()
            para.alignment = 2  # Right align
            run = para.add_run(place_date)
            run.bold = True
        title = _fill_placeholders((header.findtext("Title") or ""), ctx)
        if title:
            para = doc.add_paragraph()
            para.alignment = 1  # Center
            run = para.add_run(title)
            run.bold = True
            run.font.size = doc.styles['Heading 1'].font.size

    legal_base = root.find("LegalBase")
    if legal_base is not None:
        legal_text = _fill_placeholders((legal_base.text or ""), ctx)
        if legal_text:
            doc.add_paragraph(legal_text)

    applicants = root.find("Applicants")
    if applicants is not None:
        app_text = _fill_placeholders((applicants.text or ""), ctx)
        if app_text:
            doc.add_paragraph(app_text)

    investment = root.find("Investment")
    if investment is not None:
        inv_text = _fill_placeholders((investment.text or ""), ctx)
        if inv_text:
            doc.add_paragraph(inv_text)

    location = root.find("Location")
    if location is not None:
        loc_text = _fill_placeholders((location.text or ""), ctx)
        if loc_text:
            doc.add_paragraph(loc_text)

    decision_title = root.find("DecisionTitle")
    if decision_title is not None:
        title_text = _fill_placeholders((decision_title.text or ""), ctx)
        if title_text:
            para = doc.add_paragraph()
            para.alignment = 1  # Center
            run = para.add_run(title_text)
            run.bold = True
            run.font.size = doc.styles['Heading 1'].font.size

    decision_intro = root.find("DecisionIntro")
    if decision_intro is not None:
        intro_text = _fill_placeholders((decision_intro.text or ""), ctx)
        if intro_text:
            doc.add_paragraph(intro_text)

    conditions = root.find("Conditions")
    if conditions is not None:
        for point in conditions.findall("Point"):
            idx = point.get("index") or ""
            title = _fill_placeholders((point.get("title") or ""), ctx)
            if title:
                para = doc.add_paragraph()
                run = para.add_run(f"{idx}. {title}")
                run.bold = True
            for child in point:
                if child.tag == "Text":
                    doc.add_paragraph(_fill_placeholders((child.text or ""), ctx))
                elif child.tag == "Subpoint":
                    sidx = child.get("index") or ""
                    stitle = _fill_placeholders((child.get("title") or ""), ctx)
                    if stitle:
                        para = doc.add_paragraph()
                        run = para.add_run(f"{sidx}. {stitle}")
                        run.bold = True
                    for subchild in child:
                        if subchild.tag == "Text":
                            doc.add_paragraph(_fill_placeholders((subchild.text or ""), ctx))
                        elif subchild.tag == "Item":
                            item_idx = subchild.get("index") or ""
                            label = _fill_placeholders((subchild.findtext("Label") or ""), ctx)
                            text = _fill_placeholders((subchild.findtext("Text") or ""), ctx)
                            para = doc.add_paragraph()
                            run = para.add_run(f"{item_idx}) ")
                            run.bold = True
                            run2 = para.add_run(f"{label}: {text}")

    justification = root.find("Justification")
    if justification is not None:
        j_title = justification.find("Title")
        if j_title is not None:
            title_text = _fill_placeholders((j_title.text or ""), ctx)
            if title_text:
                para = doc.add_paragraph()
                run = para.add_run(title_text)
                run.bold = True
        for text_elem in justification.findall("Text"):
            text_content = _fill_placeholders((text_elem.text or ""), ctx)
            if text_content:
                doc.add_paragraph(text_content)

    instruction = root.find("Instruction")
    if instruction is not None:
        i_title = instruction.find("Title")
        if i_title is not None:
            title_text = _fill_placeholders((i_title.text or ""), ctx)
            if title_text:
                para = doc.add_paragraph()
                run = para.add_run(title_text)
                run.bold = True
        for text_elem in instruction.findall("Text"):
            text_content = _fill_placeholders((text_elem.text or ""), ctx)
            if text_content:
                doc.add_paragraph(text_content)

    additional_info = root.find("AdditionalInfo")
    if additional_info is not None:
        ai_title = additional_info.find("Title")
        if ai_title is not None:
            title_text = _fill_placeholders((ai_title.text or ""), ctx)
            if title_text:
                para = doc.add_paragraph()
                run = para.add_run(title_text)
                run.bold = True
        for item in additional_info.findall("Item"):
            idx = item.get("index") or ""
            text_content = _fill_placeholders((item.findtext("Text") or ""), ctx)
            if text_content:
                para = doc.add_paragraph()
                run = para.add_run(f"{idx}. {text_content}")

    agreement = root.find("Agreement")
    if agreement is not None:
        a_title = agreement.find("Title")
        if a_title is not None:
            title_text = _fill_placeholders((a_title.text or ""), ctx)
            if title_text:
                para = doc.add_paragraph()
                run = para.add_run(title_text)
                run.bold = True
        for item in agreement.findall("Item"):
            idx = item.get("index") or ""
            text_content = _fill_placeholders((item.findtext("Text") or ""), ctx)
            if text_content:
                para = doc.add_paragraph()
                run = para.add_run(f"{idx}) {text_content}")
        note = agreement.find("Note")
        if note is not None:
            note_text = _fill_placeholders((note.text or ""), ctx)
            if note_text:
                para = doc.add_paragraph()
                run = para.add_run(note_text)
                run.italic = True
        text_elem = agreement.find("Text")
        if text_elem is not None:
            text_content = _fill_placeholders((text_elem.text or ""), ctx)
            if text_content:
                doc.add_paragraph(text_content)

    annexes = root.find("Annexes")
    if annexes is not None:
        an_title = annexes.find("Title")
        if an_title is not None:
            title_text = _fill_placeholders((an_title.text or ""), ctx)
            if title_text:
                para = doc.add_paragraph()
                run = para.add_run(title_text)
                run.bold = True
        for item in annexes.findall("Item"):
            idx = item.get("index") or ""
            text_content = _fill_placeholders((item.findtext("Text") or ""), ctx)
            if text_content:
                para = doc.add_paragraph()
                run = para.add_run(f"{idx}. {text_content}")

    recipients = root.find("Recipients")
    if recipients is not None:
        r_title = recipients.find("Title")
        if r_title is not None:
            title_text = _fill_placeholders((r_title.text or ""), ctx)
            if title_text:
                para = doc.add_paragraph()
                run = para.add_run(title_text)
                run.bold = True
        for item in recipients.findall("Item"):
            idx = item.get("index") or ""
            text_content = _fill_placeholders((item.findtext("Text") or ""), ctx)
            if text_content:
                para = doc.add_paragraph()
                run = para.add_run(f"{idx}. {text_content}")

    footer = root.find("Footer")
    if footer is not None:
        doc.add_paragraph("")
        sign_line = footer.findtext("SignLine") or ""
        if sign_line:
            doc.add_paragraph(sign_line)
        sign_hint = footer.findtext("SignHint") or ""
        if sign_hint:
            doc.add_paragraph(_fill_placeholders(sign_hint, ctx))

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


def generate_decision_pdf_from_xml(analysis: Dict[str, str], wniosek: Dict[str, str], gmina: str, case_number: str) -> bytes:
    """Generate PDF decision document from XML template"""
    # Znajdź konfigurację gminy
    municipalities = load_municipality_templates()
    gmina_key = gmina.lower().replace("gmina ", "").replace(" ", "_")
    
    if gmina_key not in municipalities:
        gmina_key = "konopnica"  # fallback
    
    municipality_config = municipalities[gmina_key]
    xml_path = os.path.join(os.path.dirname(__file__), municipality_config["templates"]["decision"])
    
    if not os.path.exists(xml_path):
        raise FileNotFoundError(f"Decision XML template not found: {xml_path}")

    # Merge analysis and wniosek data for context
    # Używamy danych bezpośrednio - XML oczekuje pól z prefixem wniosek_
    ctx = {}
    ctx.update(wniosek)  # Zawiera pola z prefixem wniosek_
    ctx.update(analysis)  # Zawiera pola bez prefixu
    ctx["gmina"] = f"Gmina {gmina}"
    ctx["case_number"] = case_number
    # Fill missing date fields if not provided
    if "data" not in ctx or not ctx["data"]:
        ctx["data"] = datetime.now().strftime("%d.%m.%Y r.")
    if "data_wniosku" not in ctx or not ctx["data_wniosku"]:
        ctx["data_wniosku"] = datetime.now().strftime("%d.%m.%Y")
    if "data_uzupełnienia" not in ctx or not ctx["data_uzupełnienia"]:
        ctx["data_uzupełnienia"] = datetime.now().strftime("%d.%m.%Y")
    if "rodzaj_zabudowy" not in ctx or not ctx["rodzaj_zabudowy"]:
        ctx["rodzaj_zabudowy"] = "zabudowa mieszkaniowa jednorodzinna"

    bio = BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    family = ensure_font_installed()
    font_regular_path = os.path.join(BASE_DIR, 'static', 'fonts', 'DejaVuSans.ttf')
    font_bold_path = os.path.join(BASE_DIR, 'static', 'fonts', 'DejaVuSans-Bold.ttf')
    base_font = 'Helvetica'
    if family and os.path.exists(font_regular_path):
        try:
            pdfmetrics.registerFont(TTFont('DejaVuSans', font_regular_path))
            if os.path.exists(font_bold_path):
                pdfmetrics.registerFont(TTFont('DejaVuSans-Bold', font_bold_path))
                pdfmetrics.registerFontFamily('DejaVu', normal='DejaVuSans', bold='DejaVuSans-Bold', italic='DejaVuSans', boldItalic='DejaVuSans-Bold')
            base_font = 'DejaVu'
        except Exception:
            base_font = 'Helvetica'
    if base_font == 'DejaVu':
        styles['Normal'].fontName = 'DejaVuSans'
        if 'Title' in styles.byName:
            styles['Title'].fontName = 'DejaVuSans-Bold'
        if 'Heading2' in styles.byName:
            styles['Heading2'].fontName = 'DejaVuSans-Bold'
    elements = []

    tree = ET.parse(xml_path)
    root = tree.getroot()

    header = root.find("Header")
    if header is not None:
        ref_num = _fill_placeholders((header.findtext("ReferenceNumber") or ""), ctx)
        if ref_num:
            elements.append(Paragraph(f"<b>{ref_num}</b>", styles['Normal']))
        place_date = _fill_placeholders((header.findtext("PlaceDate") or ""), ctx)
        if place_date:
            para = Paragraph(f"<b>{place_date}</b>", styles['Normal'])
            para.alignment = 2  # Right align
            elements.append(para)
        title = _fill_placeholders((header.findtext("Title") or ""), ctx)
        if title:
            para = Paragraph(f"<b>{title}</b>", styles['Title'])
            para.alignment = 1  # Center
            elements.append(para)
        elements.append(Spacer(1, 12))

    legal_base = root.find("LegalBase")
    if legal_base is not None:
        legal_text = _fill_placeholders((legal_base.text or ""), ctx)
        if legal_text:
            elements.append(Paragraph(legal_text.replace('\n', '<br/>'), styles['Normal']))
            elements.append(Spacer(1, 6))

    applicants = root.find("Applicants")
    if applicants is not None:
        app_text = _fill_placeholders((applicants.text or ""), ctx)
        if app_text:
            elements.append(Paragraph(app_text.replace('\n', '<br/>'), styles['Normal']))
            elements.append(Spacer(1, 6))

    investment = root.find("Investment")
    if investment is not None:
        inv_text = _fill_placeholders((investment.text or ""), ctx)
        if inv_text:
            elements.append(Paragraph(inv_text.replace('\n', '<br/>'), styles['Normal']))
            elements.append(Spacer(1, 6))

    location = root.find("Location")
    if location is not None:
        loc_text = _fill_placeholders((location.text or ""), ctx)
        if loc_text:
            elements.append(Paragraph(loc_text.replace('\n', '<br/>'), styles['Normal']))
            elements.append(Spacer(1, 6))

    decision_title = root.find("DecisionTitle")
    if decision_title is not None:
        title_text = _fill_placeholders((decision_title.text or ""), ctx)
        if title_text:
            para = Paragraph(f"<b>{title_text}</b>", styles['Title'])
            para.alignment = 1  # Center
            elements.append(para)
            elements.append(Spacer(1, 8))

    decision_intro = root.find("DecisionIntro")
    if decision_intro is not None:
        intro_text = _fill_placeholders((decision_intro.text or ""), ctx)
        if intro_text:
            elements.append(Paragraph(intro_text.replace('\n', '<br/>'), styles['Normal']))
            elements.append(Spacer(1, 8))

    conditions = root.find("Conditions")
    if conditions is not None:
        for point in conditions.findall("Point"):
            idx = point.get("index") or ""
            title = _fill_placeholders((point.get("title") or ""), ctx)
            if title:
                elements.append(Paragraph(f"<b>{idx}. {title}</b>", styles['Normal']))
            for child in point:
                if child.tag == "Text":
                    elements.append(Paragraph(_fill_placeholders((child.text or ""), ctx).replace('\n', '<br/>'), styles['Normal']))
                elif child.tag == "Subpoint":
                    sidx = child.get("index") or ""
                    stitle = _fill_placeholders((child.get("title") or ""), ctx)
                    if stitle:
                        elements.append(Paragraph(f"<b>{sidx}. {stitle}</b>", styles['Normal']))
                    for subchild in child:
                        if subchild.tag == "Text":
                            elements.append(Paragraph(_fill_placeholders((subchild.text or ""), ctx).replace('\n', '<br/>'), styles['Normal']))
                        elif subchild.tag == "Item":
                            item_idx = subchild.get("index") or ""
                            label = _fill_placeholders((subchild.findtext("Label") or ""), ctx)
                            text = _fill_placeholders((subchild.findtext("Text") or ""), ctx)
                            elements.append(Paragraph(f"{item_idx}) <b>{label}:</b> {text}", styles['Normal']))
            elements.append(Spacer(1, 6))

    justification = root.find("Justification")
    if justification is not None:
        j_title = justification.find("Title")
        if j_title is not None:
            title_text = _fill_placeholders((j_title.text or ""), ctx)
            if title_text:
                elements.append(Paragraph(f"<b>{title_text}</b>", styles['Heading2']))
        for text_elem in justification.findall("Text"):
            text_content = _fill_placeholders((text_elem.text or ""), ctx)
            if text_content:
                elements.append(Paragraph(text_content.replace('\n', '<br/>'), styles['Normal']))
                elements.append(Spacer(1, 6))

    instruction = root.find("Instruction")
    if instruction is not None:
        i_title = instruction.find("Title")
        if i_title is not None:
            title_text = _fill_placeholders((i_title.text or ""), ctx)
            if title_text:
                elements.append(Paragraph(f"<b>{title_text}</b>", styles['Heading2']))
        for text_elem in instruction.findall("Text"):
            text_content = _fill_placeholders((text_elem.text or ""), ctx)
            if text_content:
                elements.append(Paragraph(text_content.replace('\n', '<br/>'), styles['Normal']))
                elements.append(Spacer(1, 6))

    additional_info = root.find("AdditionalInfo")
    if additional_info is not None:
        ai_title = additional_info.find("Title")
        if ai_title is not None:
            title_text = _fill_placeholders((ai_title.text or ""), ctx)
            if title_text:
                elements.append(Paragraph(f"<b>{title_text}</b>", styles['Heading2']))
        for item in additional_info.findall("Item"):
            idx = item.get("index") or ""
            text_content = _fill_placeholders((item.findtext("Text") or ""), ctx)
            if text_content:
                elements.append(Paragraph(f"<b>{idx}.</b> {text_content.replace('\n', '<br/>')}", styles['Normal']))
                elements.append(Spacer(1, 4))

    agreement = root.find("Agreement")
    if agreement is not None:
        a_title = agreement.find("Title")
        if a_title is not None:
            title_text = _fill_placeholders((a_title.text or ""), ctx)
            if title_text:
                elements.append(Paragraph(f"<b>{title_text}</b>", styles['Heading2']))
        for item in agreement.findall("Item"):
            idx = item.get("index") or ""
            text_content = _fill_placeholders((item.findtext("Text") or ""), ctx)
            if text_content:
                elements.append(Paragraph(f"<b>{idx})</b> {text_content.replace('\n', '<br/>')}", styles['Normal']))
                elements.append(Spacer(1, 4))
        note = agreement.find("Note")
        if note is not None:
            note_text = _fill_placeholders((note.text or ""), ctx)
            if note_text:
                elements.append(Paragraph(f"<i>{note_text}</i>", styles['Normal']))
                elements.append(Spacer(1, 4))
        text_elem = agreement.find("Text")
        if text_elem is not None:
            text_content = _fill_placeholders((text_elem.text or ""), ctx)
            if text_content:
                elements.append(Paragraph(text_content.replace('\n', '<br/>'), styles['Normal']))
                elements.append(Spacer(1, 6))

    annexes = root.find("Annexes")
    if annexes is not None:
        an_title = annexes.find("Title")
        if an_title is not None:
            title_text = _fill_placeholders((an_title.text or ""), ctx)
            if title_text:
                elements.append(Paragraph(f"<b>{title_text}</b>", styles['Heading2']))
        for item in annexes.findall("Item"):
            idx = item.get("index") or ""
            text_content = _fill_placeholders((item.findtext("Text") or ""), ctx)
            if text_content:
                elements.append(Paragraph(f"<b>{idx}.</b> {text_content.replace('\n', '<br/>')}", styles['Normal']))
                elements.append(Spacer(1, 4))

    recipients = root.find("Recipients")
    if recipients is not None:
        r_title = recipients.find("Title")
        if r_title is not None:
            title_text = _fill_placeholders((r_title.text or ""), ctx)
            if title_text:
                elements.append(Paragraph(f"<b>{title_text}</b>", styles['Heading2']))
        for item in recipients.findall("Item"):
            idx = item.get("index") or ""
            text_content = _fill_placeholders((item.findtext("Text") or ""), ctx)
            if text_content:
                elements.append(Paragraph(f"<b>{idx}.</b> {text_content.replace('\n', '<br/>')}", styles['Normal']))
                elements.append(Spacer(1, 4))

    footer = root.find("Footer")
    if footer is not None:
        elements.append(Spacer(1, 12))
        sign_line = footer.findtext("SignLine") or ""
        if sign_line:
            elements.append(Paragraph(sign_line, styles['Normal']))
        sign_hint = footer.findtext("SignHint") or ""
        if sign_hint:
            elements.append(Paragraph(_fill_placeholders(sign_hint, ctx), styles['Normal']))

    doc.build(elements)
    bio.seek(0)
    return bio.read()


def field_labels() -> Dict[str, str]:
    """Ładuje definicję pól formularza z pliku fields.json"""
    # Spróbuj załadować z pliku fields.json
    fields_file = os.path.join(os.path.dirname(__file__), "fields.json")
    if os.path.exists(fields_file):
        try:
            with open(fields_file, "r", encoding="utf-8") as f:
                data = json.load(f)
                if isinstance(data, dict):
                    return data
        except Exception:
            pass
    
    # Fallback - standardowe pola (deprecated)
    return {
        "applicant_name": "Wnioskodawca",
        "applicant_address": "Adres wnioskodawcy",
        "parcel_numbers": "Numery działek",
        "precinct": "Obręb",
        "municipality": "Gmina (z wniosku)",
        "investment_description": "Opis inwestycji",
        "land_use": "Przeznaczenie terenu",
        "building_height": "Wysokość zabudowy",
        "building_area": "Powierzchnia zabudowy",
        "parking": "Miejsca parkingowe",
        "wyniki_analizy": "Wyniki analizy (tekst)",
        "uzasadnienie": "Uzasadnienie (tekst)",
        "podstawy_prawne": "Podstawy prawne (tekst)",
    }


def long_text_keys() -> List[str]:
    return [
        "wyniki_analizy",
        "uzasadnienie",
        "podstawy_prawne",
    ]


def wniosek_only_keys() -> List[str]:
    """Pola które są wypełniane tylko w sekcji wniosku i automatycznie kopiowane do analizy"""
    return [
        "wnioskodawca_mianownik",
        "wnioskodawca_dopelniacz",
        "wnioskodawca_adres",
        "gmina",
        "obreb",
        "dzialki",
        "teren_obejmuje",
        "data_wykonania_analizy",
        "data_zlozenia_wniosku",
        "data_uzupelnienia_wniosku",
        "rodzaj_inwestycji",
        "opis_inwestycji",
        "linia_zabudowy",
        "szerokosc_elewacji_frontowej",
        "wysokosc_zabudowy",
        "liczba_kondygnacji",
        "kat_nachylenia_dachu",
        "powierzchnia_zabudowy",
        "powierzchnia_biologicznie_czynna",
        "intensywnosc_zabudowy",
        "miejsca_parkingowe",
        "dostep_droga_publiczna",
        "woda",
        "scieki",
        "odwodnienie",
        "energia_elektr",
        "gaz",
        "ogrzewanie",
        "odpady",
        "uwagi",
        "data",
        "podpis",
        "rodzaj_zabudowy"
    ]


def required_fields() -> List[str]:
    """Pola wymagane do wypełnienia - tylko te które użytkownik wskaże"""
    return [
        "wnioskodawca_mianownik",
        "wnioskodawca_dopelniacz",
        "wnioskodawca_adres",
        "gmina",
        "obreb",
        "dzialki",
        "data_wykonania_analizy",
        "data_zlozenia_wniosku",
    ]


def process_dzialki_fields(form) -> Dict[str, str]:
    """Zbiera wszystkie pola działek i zwraca flagę oraz połączoną wartość"""
    dzialki_values = []
    
    # Najpierw sprawdź pole bez indeksu (pierwsze pole)
    val = str(form.get("dzialki_wniosek", "")).strip()
    if val:
        dzialki_values.append(val)
    
    # Następnie sprawdź pola z indeksem
    i = 1
    while True:
        val = str(form.get(f"dzialki_wniosek_{i}", "")).strip()
        if val:
            dzialki_values.append(val)
        i += 1
        
        # Zatrzymaj gdy brak kolejnego pola (sprawdź dla FastAPI Form)
        if hasattr(form, 'getlist'):
            # FastAPI Form object
            if not form.get(f"dzialki_wniosek_{i}", ""):
                break
        else:
            # Dictionary
            if f"dzialki_wniosek_{i}" not in form:
                break
    
    # Połącz wartości
    dzialki_combined = ", ".join(dzialki_values) if dzialki_values else ""
    
    # Określ flagę
    is_multiple = len(dzialki_values) > 1
    
    return {
        "dzialki": dzialki_combined,
        "dzialki_multiple": "true" if is_multiple else "false",
        "dzialki_count": str(len(dzialki_values))
    }

def process_data_uzupelnienia_fields(form) -> str:
    """Zbiera wszystkie daty uzupełnienia i zwraca połączoną wartość"""
    dates_values = []
    
    # Najpierw sprawdź pole bez indeksu (pierwsze pole)
    val = str(form.get("data_uzupelnienia_wniosku_wniosek", "")).strip()
    if val:
        dates_values.append(val)
    
    # Następnie sprawdź pola z indeksem
    i = 1
    while True:
        val = str(form.get(f"data_uzupelnienia_wniosku_wniosek_{i}", "")).strip()
        if val:
            dates_values.append(val)
        i += 1
        
        # Zatrzymaj gdy brak kolejnego pola (sprawdź dla FastAPI Form)
        if hasattr(form, 'getlist'):
            # FastAPI Form object
            if not form.get(f"data_uzupelnienia_wniosku_wniosek_{i}", ""):
                break
        else:
            # Dictionary
            if f"data_uzupelnienia_wniosku_wniosek_{i}" not in form:
                break
    
    # Połącz wartości (daty będą już w formacie DD.MM.YYYY z JavaScript)
    return ", ".join(dates_values) if dates_values else ""


def transform_title_to_dopelniacz(title: str) -> str:
    """Przekształca tytuł na dopełniacz zgodnie z regułami"""
    if not title:
        return ""
    
    # Słownik przekształceń tytułów
    title_transformations = {
        "Pan": "Pana",
        "Pani": "Pani", 
        "Państwo": "Państwa",
        "Podmiot": "Podmiotu"
    }
    
    return title_transformations.get(title, title)

def transform_mianownik_to_dopelniacz(mianownik: str) -> str:
    """Przekształca mianownik na dopełniacz zgodnie z regułami"""
    if not mianownik:
        return ""
    
    # Słownik przekształceń tytułów
    title_transformations = {
        "Pan": "Pana",
        "Pani": "Pani", 
        "Państwo": "Państwa",
        "Podmiot": "Podmiotu"
    }
    
    # Sprawdź czy wartość zaczyna się od tytułu
    for title, dopelniacz_title in title_transformations.items():
        if mianownik.startswith(f"{title} "):
            # Zamień tytuł na dopełniacz i zwróć resztę bez zmian
            return mianownik.replace(f"{title} ", f"{dopelniacz_title} ", 1)
    
    # Jeśli nie ma tytułu, zwróć bez zmian
    return mianownik


def validate_required_fields(form_data: Dict[str, str]) -> List[str]:
    """Waliduje wymagane pola i zwraca listę błędów"""
    errors = []
    required = required_fields()
    
    for field in required:
        value = form_data.get(field, "").strip()
        if not value:
            field_label = field_labels().get(field, field)
            errors.append(f"Pole '{field_label}' jest wymagane")
    
    # Specjalna walidacja dla wnioskodawca_mianownik - sprawdź czy wybrano tytuł
    wnioskodawca_value = form_data.get("wnioskodawca_mianownik", "").strip()
    if wnioskodawca_value:
        # Sprawdź czy wartość zaczyna się od tytułu (Pan, Pani, Państwo, Podmiot)
        if not any(wnioskodawca_value.startswith(prefix) for prefix in ["Pan", "Pani", "Państwo", "Podmiot"]):
            errors.append("Dla pola 'Wnioskodawca - Mianownik' należy wybrać tytuł (Pan/Pani/Państwo/Podmiot)")
    
    # Walidacja dat - sprawdź czy data złożenia wniosku nie jest późniejsza niż data analizy
    data_zlozenia = form_data.get("data_zlozenia_wniosku_wniosek", "").strip()
    data_analizy = form_data.get("data_wykonania_analizy_wniosek", "").strip()
    
    if data_zlozenia and data_analizy:
        try:
            # Parsuj daty YYYY-MM-DD
            from datetime import datetime
            date_zlozenia = datetime.strptime(data_zlozenia, "%Y-%m-%d")
            date_analizy = datetime.strptime(data_analizy, "%Y-%m-%d")
            
            if date_zlozenia > date_analizy:
                errors.append("Data złożenia wniosku nie może być późniejsza niż data wykonania analizy")
        except ValueError:
            # Błąd parsowania - ignoruj, bo walidacja formatu zostanie obsłużona gdzie indziej
            pass
    
    # Walidacja dat uzupełnienia wniosku
    from datetime import datetime
    
    # Zbierz wszystkie daty uzupełnienia
    daty_uzupelnienia = []
    i = 0
    while True:
        key = f"data_uzupelnienia_wniosku_wniosek_{i}" if i > 0 else "data_uzupelnienia_wniosku_wniosek"
        val = form_data.get(key, "").strip()
        if val:
            try:
                # Parsuj YYYY-MM-DD lub DD.MM.YYYY
                if '.' in val:
                    date_obj = datetime.strptime(val, "%d.%m.%Y")
                else:
                    date_obj = datetime.strptime(val, "%Y-%m-%d")
                daty_uzupelnienia.append((key, date_obj))
            except ValueError:
                pass
        i += 1
        if i > 100:  # Bezpieczny limit
            break
    
    # Waliduj daty uzupełnienia
    if daty_uzupelnienia and data_zlozenia and data_analizy:
        try:
            date_zlozenia = datetime.strptime(data_zlozenia, "%Y-%m-%d")
            date_analizy = datetime.strptime(data_analizy, "%Y-%m-%d")
            
            for i, (key, date_uz) in enumerate(daty_uzupelnienia):
                # 1. Data nie może być wcześniejsza niż data złożenia wniosku
                if date_uz < date_zlozenia:
                    errors.append(f"Data uzupełnienia wniosku nie może być wcześniejsza niż data złożenia wniosku")
                
                # 2. Data nie może być późniejsza niż data wykonania analizy
                if date_uz > date_analizy:
                    errors.append(f"Data uzupełnienia wniosku nie może być późniejsza niż data wykonania analizy")
                
                # 3. Kolejna data nie może być wcześniejsza niż poprzednia
                if i > 0:
                    prev_date = daty_uzupelnienia[i-1][1]
                    if date_uz < prev_date:
                        errors.append(f"Data uzupełnienia wniosku nie może być wcześniejsza niż poprzednia data")
        except ValueError:
            pass
    
    return errors


def validate_radio_buttons(form_data) -> List[str]:
    """Waliduje radio buttony - wymaga wyboru jednej opcji"""
    errors = []
    
    # Sprawdź czy wybrano tytuł dla wnioskodawcy
    # Obsługuj zarówno słowniki jak i obiekty FastAPI Form
    if hasattr(form_data, 'get'):
        # FastAPI Form object - użyj get() żeby wziąć pierwszą wartość dla klucza
        title_val = str(form_data.get("wnioskodawca_title_wniosek", "")).strip()
    else:
        # Dictionary
        title_val = str(form_data.get("wnioskodawca_title_wniosek", "")).strip()
    
    title_selected = title_val in ["Pan", "Pani", "Państwo", "Podmiot"]
    
    if not title_selected:
        errors.append("Należy wybrać tytuł dla pola 'Wnioskodawca - Mianownik' (Pan/Pani/Państwo/Podmiot)")
    
    return errors


app = FastAPI(title="Analiza urbanistyczna - WZ")

BASE_DIR = os.path.dirname(os.path.dirname(__file__))
templates = Jinja2Templates(directory=os.path.join(BASE_DIR, "templates"))
static_dir = os.path.join(BASE_DIR, "static")
if not os.path.exists(static_dir):
    os.makedirs(static_dir, exist_ok=True)
app.mount("/static", StaticFiles(directory=static_dir), name="static")


def ensure_font_installed() -> str:
    """Ensure DejaVuSans (regular/bold) exists locally; download if missing. Returns base family name or empty."""
    fonts_dir = os.path.join(static_dir, "fonts")
    os.makedirs(fonts_dir, exist_ok=True)
    regular = os.path.join(fonts_dir, "DejaVuSans.ttf")
    bold = os.path.join(fonts_dir, "DejaVuSans-Bold.ttf")
    if not (os.path.exists(regular) and os.path.getsize(regular) > 0):
        try:
            urllib.request.urlretrieve(
                "https://github.com/dejavu-fonts/dejavu-fonts/raw/version_2_37/ttf/DejaVuSans.ttf",
                regular,
            )
        except Exception:
            pass
    if not (os.path.exists(bold) and os.path.getsize(bold) > 0):
        try:
            urllib.request.urlretrieve(
                "https://github.com/dejavu-fonts/dejavu-fonts/raw/version_2_37/ttf/DejaVuSans-Bold.ttf",
                bold,
            )
        except Exception:
            pass
    # Return family name if at least regular exists
    return "DejaVu" if os.path.exists(regular) else ""


# Prepare font at startup (best-effort)
ensure_font_installed()


def extract_pdf_form_fields(file_path: str) -> Dict[str, str]:
    reader = PdfReader(file_path)
    fields = reader.get_fields() or {}
    mapping: Dict[str, str] = {}
    for key, field in fields.items():
        name = str(key)
        mapping[name] = name
    return mapping


# Static fields are loaded from app/fields.json if present; otherwise fallback to code defaults


@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    municipalities = load_municipality_templates()
    # Wyciągnij nazwy gmin z konfiguracji (bez prefiksu "Gmina")
    gminas = [config['name'] for config in municipalities.values()]
    return templates.TemplateResponse(
        "index.html",
        {
            "request": request,
            "gminas": gminas,
            "labels": field_labels(),
            "case_number": "",
            "long_text_keys": long_text_keys(),
            "wniosek_only_keys": wniosek_only_keys(),
        },
    )


@app.post("/compare", response_class=HTMLResponse)
async def compare(request: Request, gmina: str = Form(...), case_number: str = Form("")):
    form = await request.form()
    labels = field_labels()
    wniosek: Dict[str, str] = {}
    analiza: Dict[str, str] = {}
    wniosek_only = wniosek_only_keys()
    
    # Obsługa specjalna dla działek (wiele pól)
    dzialki_data = process_dzialki_fields(form)
    
    # Obsługa specjalna dla dat uzupełnienia (wiele pól)
    data_uzupelnienia_combined = process_data_uzupelnienia_fields(form)
    
    for key in labels.keys():
        # Pomiń działki i data_uzupelnienia_wniosku - będą przetworzone specjalnie
        if key == "dzialki":
            continue
        if key == "data_uzupelnienia_wniosku":
            continue
            
        val = str(form.get(f"{key}_wniosek", ""))
        
        # Obsługa radio buttonów dla wnioskodawca_mianownik
        if key == "wnioskodawca_mianownik":
            title_val = str(form.get("wnioskodawca_title_wniosek", ""))
            if title_val:
                val = f"{title_val} {val}"
        
        # Obsługa radio buttonów dla wnioskodawca_dopelniacz
        elif key == "wnioskodawca_dopelniacz":
            title_val = str(form.get("wnioskodawca_title_wniosek", ""))
            if title_val:
                # Przekształć tytuł na dopełniacz
                dopelniacz_title = transform_title_to_dopelniacz(title_val)
                val = f"{dopelniacz_title} {val}"
        
        wniosek[f"{key}"] = val
        
        # Dla pól "tylko wniosek" - kopiuj z wniosku do analizy
        if key in wniosek_only:
            analiza[f"{key}"] = val
        else:
            analiza[f"{key}"] = str(form.get(f"{key}_analiza", ""))
    
    # Dodaj przetworzone dane działek (bez prefixu wniosek_ dla /compare)
    wniosek["dzialki"] = dzialki_data["dzialki"]
    wniosek["dzialki_multiple"] = dzialki_data["dzialki_multiple"]
    wniosek["dzialki_count"] = dzialki_data["dzialki_count"]
    analiza["dzialki"] = dzialki_data["dzialki"]
    analiza["dzialki_multiple"] = dzialki_data["dzialki_multiple"]
    analiza["dzialki_count"] = dzialki_data["dzialki_count"]
    
    # Dodaj przetworzone daty uzupełnienia
    wniosek["data_uzupelnienia_wniosku"] = data_uzupelnienia_combined
    analiza["data_uzupelnienia_wniosku"] = data_uzupelnienia_combined

    # Walidacja wymaganych pól i radio buttonów
    validation_errors = validate_required_fields(wniosek)
    radio_errors = validate_radio_buttons(form)
    validation_errors.extend(radio_errors)
    
    comparison = compare_values(wniosek, analiza)
    municipalities = load_municipality_templates()
    gminas = [config['name'] for config in municipalities.values()]
    return templates.TemplateResponse(
        "index.html",
        {
            "request": request,
            "gminas": gminas,
            "labels": field_labels(),
            "selected_gmina": gmina,
            "comparison": comparison,
            "case_number": case_number,
            "long_text_keys": long_text_keys(),
            "wniosek_only_keys": wniosek_only_keys(),
            "validation_errors": validation_errors,
        },
    )


@app.post("/generate-docx")
async def generate_docx_endpoint(request: Request, gmina: str = Form(...), case_number: str = Form("")):
    form = await request.form()
    labels = field_labels()
    wniosek_only = wniosek_only_keys()
    
    # Zbierz dane z formularza
    wniosek = {}
    analysis = {}
    
    for key in labels.keys():
        val = str(form.get(f"{key}_wniosek", "")).strip()
        
        # Obsługa radio buttonów dla wnioskodawca_mianownik
        if key == "wnioskodawca_mianownik":
            # Użyj get() zamiast getlist() - zwróci tylko jedną wartość dla klucza
            # To zapewni, że weźmiemy wartość z sekcji wniosek, nie z analizy
            title_val = str(form.get("wnioskodawca_title_wniosek", "")).strip()
            # Dodaj tytuł tylko jeśli pole tekstowe nie jest puste
            if title_val and val:
                val = f"{title_val} {val}"
            elif title_val and not val:
                # Tylko tytuł bez wartości - zostaw puste (walidacja to złapie)
                val = ""
        
        # Obsługa radio buttonów dla wnioskodawca_dopelniacz
        elif key == "wnioskodawca_dopelniacz":
            title_val = str(form.get("wnioskodawca_title_wniosek", "")).strip()
            if title_val and val:
                # Przekształć tytuł na dopełniacz
                dopelniacz_title = transform_title_to_dopelniacz(title_val)
                val = f"{dopelniacz_title} {val}"
            elif title_val and not val:
                # Tylko tytuł bez wartości - zostaw puste (walidacja to złapie)
                val = ""
        
        if key in wniosek_only:
            # Dla pól "tylko wniosek" - użyj przedrostka wniosek_
            wniosek[f"wniosek_{key}"] = val
            analysis[f"wniosek_{key}"] = val
        else:
            wniosek[key] = val
            analysis[key] = str(form.get(f"{key}_analiza", "")).strip()

    # Walidacja wymaganych pól i radio buttonów
    # Dla walidacji używamy danych bez prefiksu wniosek_
    wniosek_for_validation = {k.replace("wniosek_", ""): v for k, v in wniosek.items() if k.startswith("wniosek_")}
    wniosek_for_validation.update({k: v for k, v in wniosek.items() if not k.startswith("wniosek_")})
    validation_errors = validate_required_fields(wniosek_for_validation)
    radio_errors = validate_radio_buttons(form)
    validation_errors.extend(radio_errors)
    if validation_errors:
        error_html = f"<h1>Błędy walidacji:</h1><ul>{''.join(f'<li>{error}</li>' for error in validation_errors)}</ul>"
        return HTMLResponse(content=error_html, status_code=400)
    
    # Generate DOCX from XML template
    data = generate_docx_from_xml(analysis, wniosek, gmina, case_number)
    base = sanitize_case_number(case_number) or f"analiza_urbanistyczna_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    filename = f"{base}.docx"
    return StreamingResponse(
        BytesIO(data),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@app.post("/generate-pdf")
async def generate_pdf_endpoint(request: Request, gmina: str = Form(...), case_number: str = Form("")):
    form = await request.form()
    labels = field_labels()
    wniosek_only = wniosek_only_keys()
    
    # DEBUG: Sprawdź wszystkie dane z formularza
    print("DEBUG generate-pdf: All form data:")
    for key, value in form.items():
        print(f"  {key} = '{value}'")
    
    # Zbierz dane z formularza
    wniosek = {}
    analysis = {}
    
    for key in labels.keys():
        val = str(form.get(f"{key}_wniosek", "")).strip()
        
        # Obsługa radio buttonów dla wnioskodawca_mianownik
        if key == "wnioskodawca_mianownik":
            # Użyj get() zamiast getlist() - zwróci tylko jedną wartość dla klucza
            # To zapewni, że weźmiemy wartość z sekcji wniosek, nie z analizy
            title_val = str(form.get("wnioskodawca_title_wniosek", "")).strip()
            print(f"DEBUG generate-pdf: key={key}, title_val={title_val}, val={val}")
            # Dodaj tytuł tylko jeśli pole tekstowe nie jest puste
            if title_val and val:
                val = f"{title_val} {val}"
            elif title_val and not val:
                # Tylko tytuł bez wartości - zostaw puste (walidacja to złapie)
                val = ""
        
        # Obsługa radio buttonów dla wnioskodawca_dopelniacz
        elif key == "wnioskodawca_dopelniacz":
            title_val = str(form.get("wnioskodawca_title_wniosek", "")).strip()
            if title_val and val:
                # Przekształć tytuł na dopełniacz
                dopelniacz_title = transform_title_to_dopelniacz(title_val)
                val = f"{dopelniacz_title} {val}"
            elif title_val and not val:
                # Tylko tytuł bez wartości - zostaw puste (walidacja to złapie)
                val = ""
        
        if key in wniosek_only:
            # Dla pól "tylko wniosek" - użyj przedrostka wniosek_
            wniosek[f"wniosek_{key}"] = val
            analysis[f"wniosek_{key}"] = val
        else:
            wniosek[key] = val
            analysis[key] = str(form.get(f"{key}_analiza", "")).strip()

    # Walidacja wymaganych pól i radio buttonów
    # Dla walidacji używamy danych bez prefiksu wniosek_
    wniosek_for_validation = {k.replace("wniosek_", ""): v for k, v in wniosek.items() if k.startswith("wniosek_")}
    wniosek_for_validation.update({k: v for k, v in wniosek.items() if not k.startswith("wniosek_")})
    validation_errors = validate_required_fields(wniosek_for_validation)
    radio_errors = validate_radio_buttons(form)
    validation_errors.extend(radio_errors)
    if validation_errors:
        error_html = f"<h1>Błędy walidacji:</h1><ul>{''.join(f'<li>{error}</li>' for error in validation_errors)}</ul>"
        return HTMLResponse(content=error_html, status_code=400)
    
    templates_cfg = load_municipality_templates()
    # First, try via Word (docx2pdf). Fallback to pure Python PDF.
    try:
        # Try via Word first using DOCX rendered from XML
        docx_bytes = generate_docx_from_xml(analysis, wniosek, gmina, case_number)
        pdf_bytes = generate_pdf_from_docx_bytes(docx_bytes)
    except Exception:
        pdf_bytes = generate_pdf_from_xml(analysis, wniosek, gmina, case_number)
    base = sanitize_case_number(case_number) or f"analiza_urbanistyczna_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    filename = f"{base}.pdf"
    return StreamingResponse(
        BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@app.post("/generate-decision-docx")
async def generate_decision_docx_endpoint(request: Request, gmina: str = Form(...), case_number: str = Form("")):
    form = await request.form()
    labels = field_labels()
    wniosek_only = wniosek_only_keys()
    
    # DEBUG: Sprawdź wszystkie dane z formularza
    print("DEBUG: All form data:")
    for key, value in form.items():
        print(f"  {key} = '{value}'")
    
    wniosek = {}
    analiza = {}
    
    for key in labels.keys():
        if key in wniosek_only:
            # Dla pól "tylko wniosek" - użyj przedrostka wniosek_
            val = str(form.get(f"{key}_wniosek", "")).strip()
            # Obsługa radio buttonów dla wnioskodawca_mianownik
            if key == "wnioskodawca_mianownik":
                # Użyj get() zamiast getlist() - zwróci tylko jedną wartość dla klucza
                # To zapewni, że weźmiemy wartość z sekcji wniosek, nie z analizy
                title_val = str(form.get("wnioskodawca_title_wniosek", "")).strip()
                print(f"DEBUG: {key} - title_val = '{title_val}', val = '{val}'")
                # Dodaj tytuł tylko jeśli pole tekstowe nie jest puste
                if title_val and val:
                    val = f"{title_val} {val}"
                elif title_val and not val:
                    # Tylko tytuł bez wartości - zostaw puste (walidacja to złapie)
                    val = ""
                print(f"DEBUG: {key} - final val = '{val}'")
            # Obsługa radio buttonów dla wnioskodawca_dopelniacz
            elif key == "wnioskodawca_dopelniacz":
                title_val = str(form.get("wnioskodawca_title_wniosek", "")).strip()
                print(f"DEBUG: {key} - title_val = '{title_val}', val = '{val}'")
                if title_val and val:
                    # Przekształć tytuł na dopełniacz
                    dopelniacz_title = transform_title_to_dopelniacz(title_val)
                    val = f"{dopelniacz_title} {val}"
                elif title_val and not val:
                    # Tylko tytuł bez wartości - zostaw puste (walidacja to złapie)
                    val = ""
                print(f"DEBUG: {key} - final val = '{val}'")
            wniosek[f"wniosek_{key}"] = val
            analiza[f"wniosek_{key}"] = val
        else:
            wniosek[key] = str(form.get(f"{key}_wniosek", "")).strip()
            analiza[key] = str(form.get(f"{key}_analiza", "")).strip()

    # Walidacja wymaganych pól (sprawdź wniosek bez przedrostka)
    wniosek_for_validation = {k.replace("wniosek_", ""): v for k, v in wniosek.items() if k.startswith("wniosek_")}
    wniosek_for_validation.update({k: v for k, v in wniosek.items() if not k.startswith("wniosek_")})
    validation_errors = validate_required_fields(wniosek_for_validation)
    radio_errors = validate_radio_buttons(form)
    validation_errors.extend(radio_errors)
    if validation_errors:
        error_html = f"<h1>Błędy walidacji:</h1><ul>{''.join(f'<li>{error}</li>' for error in validation_errors)}</ul>"
        return HTMLResponse(content=error_html, status_code=400)
    
    # Generate decision DOCX from XML template
    data = generate_decision_docx_from_xml(analiza, wniosek, gmina, case_number)
    
    base = sanitize_case_number(case_number) or f"decyzja_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    filename = f"{base}.docx"
    return StreamingResponse(
        BytesIO(data),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@app.post("/generate-decision-pdf")
async def generate_decision_pdf_endpoint(request: Request, gmina: str = Form(...), case_number: str = Form("")):
    form = await request.form()
    labels = field_labels()
    wniosek_only = wniosek_only_keys()
    
    wniosek = {}
    analiza = {}
    
    for key in labels.keys():
        if key in wniosek_only:
            # Dla pól "tylko wniosek" - użyj przedrostka wniosek_
            val = str(form.get(f"{key}_wniosek", "")).strip()
            # Obsługa radio buttonów dla wnioskodawca_mianownik
            if key == "wnioskodawca_mianownik":
                # Użyj get() zamiast getlist() - zwróci tylko jedną wartość dla klucza
                # To zapewni, że weźmiemy wartość z sekcji wniosek, nie z analizy
                title_val = str(form.get("wnioskodawca_title_wniosek", "")).strip()
                print(f"DEBUG: {key} - title_val = '{title_val}', val = '{val}'")
                # Dodaj tytuł tylko jeśli pole tekstowe nie jest puste
                if title_val and val:
                    val = f"{title_val} {val}"
                elif title_val and not val:
                    # Tylko tytuł bez wartości - zostaw puste (walidacja to złapie)
                    val = ""
                print(f"DEBUG: {key} - final val = '{val}'")
            # Obsługa radio buttonów dla wnioskodawca_dopelniacz
            elif key == "wnioskodawca_dopelniacz":
                title_val = str(form.get("wnioskodawca_title_wniosek", "")).strip()
                print(f"DEBUG: {key} - title_val = '{title_val}', val = '{val}'")
                if title_val and val:
                    # Przekształć tytuł na dopełniacz
                    dopelniacz_title = transform_title_to_dopelniacz(title_val)
                    val = f"{dopelniacz_title} {val}"
                elif title_val and not val:
                    # Tylko tytuł bez wartości - zostaw puste (walidacja to złapie)
                    val = ""
                print(f"DEBUG: {key} - final val = '{val}'")
            wniosek[f"wniosek_{key}"] = val
            analiza[f"wniosek_{key}"] = val
        else:
            wniosek[key] = str(form.get(f"{key}_wniosek", "")).strip()
            analiza[key] = str(form.get(f"{key}_analiza", "")).strip()

    # Walidacja wymaganych pól (sprawdź wniosek bez przedrostka)
    wniosek_for_validation = {k.replace("wniosek_", ""): v for k, v in wniosek.items() if k.startswith("wniosek_")}
    wniosek_for_validation.update({k: v for k, v in wniosek.items() if not k.startswith("wniosek_")})
    
    validation_errors = validate_required_fields(wniosek_for_validation)
    radio_errors = validate_radio_buttons(form)
    validation_errors.extend(radio_errors)
    
    if validation_errors:
        error_html = f"<h1>Błędy walidacji:</h1><ul>{''.join(f'<li>{error}</li>' for error in validation_errors)}</ul>"
        return HTMLResponse(content=error_html, status_code=400)
    
    # Try via Word first, fallback to ReportLab
    try:
        docx_bytes = generate_decision_docx_from_xml(analiza, wniosek, gmina, case_number)
        pdf_bytes = generate_pdf_from_docx_bytes(docx_bytes)
    except Exception:
        pdf_bytes = generate_decision_pdf_from_xml(analiza, wniosek, gmina, case_number)
    
    base = sanitize_case_number(case_number) or f"decyzja_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    filename = f"{base}.pdf"
    return StreamingResponse(
        BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@app.post("/save-case")
async def save_case(case_json: str = Form(...), filename: str = Form("")):
    try:
        # Validate JSON
        obj = json.loads(case_json)
        data = json.dumps(obj, ensure_ascii=False, indent=2)
        
        # Walidacja (bez blokowania zapisu) - tylko logowanie
        wniosek = obj.get("wniosek", {})
        validation_errors = validate_required_fields(wniosek)
        # Sprawdź radio buttony
        form_data = {}
        for key in wniosek.keys():
            if key == "wniosek_wnioskodawca_mianownik":
                value = wniosek.get(key, "")
                if value.startswith("Pan "):
                    form_data["wnioskodawca_title_wniosek"] = "Pan"
                elif value.startswith("Pani "):
                    form_data["wnioskodawca_title_wniosek"] = "Pani"
                elif value.startswith("Państwo "):
                    form_data["wnioskodawca_title_wniosek"] = "Państwo"
                elif value.startswith("Podmiot "):
                    form_data["wnioskodawca_title_wniosek"] = "Podmiot"
        radio_errors = validate_radio_buttons(form_data)
        validation_errors.extend(radio_errors)
        
        # Loguj błędy walidacji (bez blokowania)
        if validation_errors:
            print(f"Błędy walidacji przy zapisie: {validation_errors}")
            
    except Exception:
        # Pass through raw string if not valid JSON
        data = case_json
        
    if not filename:
        try:
            num = json.loads(case_json).get("case_number", "")
        except Exception:
            num = ""
        base = sanitize_case_number(num) or "sprawa_WZ"
        filename = f"{base}.json"
    return StreamingResponse(
        BytesIO(data.encode("utf-8")),
        media_type="application/json",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@app.post("/load-case", response_class=HTMLResponse)
async def load_case(request: Request, file: UploadFile = File(...)):
    content = await file.read()
    try:
        payload = json.loads(content.decode("utf-8"))
    except Exception:
        payload = {}

    gmina = payload.get("gmina")
    wniosek = payload.get("wniosek") or payload.get("left") or {}
    analiza = payload.get("analiza") or payload.get("right") or {}
    case_number = payload.get("case_number", "")

    labels = field_labels()
    wniosek_only = wniosek_only_keys()
    
    # Mapuj pola z przedrostkiem wniosek_ z powrotem na normalne nazwy
    normalized_wniosek = {}
    normalized_analiza = {}
    
    for key in labels.keys():
        if key in wniosek_only:
            # Dla pól "tylko wniosek" - sprawdź z przedrostkiem
            prefixed_key = f"wniosek_{key}"
            normalized_wniosek[key] = wniosek.get(prefixed_key, wniosek.get(key, ""))
            normalized_analiza[key] = analiza.get(prefixed_key, analiza.get(key, ""))
        else:
            normalized_wniosek[key] = wniosek.get(key, "")
            normalized_analiza[key] = analiza.get(key, "")
    
    # Normalize: ensure keys exist
    comparison = compare_values(normalized_wniosek, normalized_analiza)
    
    # Walidacja (bez blokowania wczytywania)
    validation_errors = validate_required_fields(normalized_wniosek)
    # Sprawdź radio buttony - symuluj form data
    form_data = {}
    for key in labels.keys():
        if key == "wnioskodawca_mianownik":
            value = normalized_wniosek.get(key, "")
            if value.startswith("Pan "):
                form_data["wnioskodawca_title_wniosek"] = "Pan"
            elif value.startswith("Pani "):
                form_data["wnioskodawca_title_wniosek"] = "Pani"
            elif value.startswith("Państwo "):
                form_data["wnioskodawca_title_wniosek"] = "Państwo"
            elif value.startswith("Podmiot "):
                form_data["wnioskodawca_title_wniosek"] = "Podmiot"
    radio_errors = validate_radio_buttons(form_data)
    validation_errors.extend(radio_errors)

    municipalities = load_municipality_templates()
    gminas = [config['name'] for config in municipalities.values()]
    return templates.TemplateResponse(
        "index.html",
        {
            "request": request,
            "gminas": gminas,
            "labels": labels,
            "selected_gmina": gmina or (gminas[0] if gminas else None),
            "comparison": comparison,
            "case_number": case_number,
            "long_text_keys": long_text_keys(),
            "wniosek_only_keys": wniosek_only_keys(),
            "validation_errors": validation_errors,
        },
    )


